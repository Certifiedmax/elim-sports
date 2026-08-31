try:
    import docx  # type: ignore
    from docx.shared import Inches, Pt, RGBColor  # type: ignore
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
except ImportError:
    print("Error: 'python-docx' package is not installed.")
    print("Install it using: pip install python-docx")
    exit(1)

doc = docx.Document()

# Adjust margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Title
title = doc.add_heading("Elim Sports E-Commerce & Retail Platform", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph("Technical Documentation & System Manual")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.italic = True
subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()

# 1. Executive Summary
doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph(
    "Elim Sports is a specialized sports retail and badminton equipment platform engineered "
    "to streamline catalog discovery, inventory control, and localized order fulfillment. "
    "Designed with a mobile-first architecture, the application bridges digital product browsing "
    "with direct conversational commerce, catering specifically to regional consumer habits where "
    "traditional card payment gateways present high checkout friction."
)

# 2. Tech Stack
doc.add_heading("2. Architecture & Technology Stack", level=1)
tech_items = [
    ("Frontend Framework", "Next.js (App Router), React, TypeScript"),
    ("Styling & Animations", "Tailwind CSS, Lucide Icons, Custom Marquee Keyframes"),
    ("Backend & Database", "Supabase (PostgreSQL with Row Level Security)"),
    ("Asset Storage", "Supabase Storage ('product-images' public bucket)"),
    ("Fulfillment API", "WhatsApp Deep Link / URI Conversational Checkout"),
    ("Hosting & CI/CD", "Vercel Edge Network with GitHub Continuous Deployment"),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Layer / Component"
hdr_cells[1].text = "Technology / Service"

for layer, tech in tech_items:
    row_cells = table.add_row().cells
    row_cells[0].text = layer
    row_cells[1].text = tech

doc.add_paragraph()

# 3. Database Schema
doc.add_heading("3. Database Schema & Data Models", level=1)
doc.add_heading("Table: public.products", level=2)
doc.add_paragraph("Stores equipment specifications, dynamic stock counts, pricing, and promotional discount values.")

schema_items = [
    ("id", "UUID", "Primary Key, Default gen_random_uuid()"),
    ("name", "TEXT", "Brand and product name"),
    ("category", "TEXT", "Badminton | Footwear | Apparel | Accessories"),
    ("price", "NUMERIC", "Active sale price in KSH"),
    ("original_price", "NUMERIC", "Base regular price for discount calculation"),
    ("stock_quantity", "INTEGER", "Available inventory units"),
    ("in_stock", "BOOLEAN", "Availability status flag"),
    ("image_url", "TEXT", "Public CDN asset link"),
    ("description", "TEXT", "Extended equipment specifications"),
]

s_table = doc.add_table(rows=1, cols=3)
s_table.alignment = WD_TABLE_ALIGNMENT.CENTER
s_hdr = s_table.rows[0].cells
s_hdr[0].text = "Column"
s_hdr[1].text = "Type"
s_hdr[2].text = "Description / Constraint"

for col, ctype, desc in schema_items:
    r_cells = s_table.add_row().cells
    r_cells[0].text = col
    r_cells[1].text = ctype
    r_cells[2].text = desc

doc.add_paragraph()

# 4. Core Features
doc.add_heading("4. Key Features & Implementation", level=1)
doc.add_paragraph("• Dynamic Discount Engine: Automatically computes discount percentage ((-% OFF) badge) when original_price > price.")
doc.add_paragraph("• Live Announcement Marquee: Reads live promo ticker text from store_settings table and scrolls across hero section.")
doc.add_paragraph("• Conversational WhatsApp Checkout: Formats cart orders and delivery info into a pre-structured, itemized receipt.")
doc.add_paragraph("• PIN-Secured Admin Portal: Allows stock updates, camera photo uploads to Supabase storage, and live ticker editing.")

# 5. Presentation Playbook
doc.add_heading("5. Project Presentation Playbook", level=1)
doc.add_paragraph("1. Problem Statement: Explain how high drop-off rates on card gateways inspired a WhatsApp conversational checkout.")
doc.add_paragraph("2. Storefront Demo: Showcase live search, category filtering, discount calculations, and the moving promo ticker.")
doc.add_paragraph("3. Checkout Flow: Add items to cart and show how WhatsApp receives a formatted receipt.")
doc.add_paragraph("4. Admin Sync: Demonstrate changing stock from the PIN-protected dashboard and seeing storefront updates in real time.")

# Save file
file_name = "Elim_Sports_Documentation.docx"
doc.save(file_name)
print(f"Successfully generated: {file_name}")
